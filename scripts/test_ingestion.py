"""
Layer 1 smoke test — run this to verify your ingestion stack works.

Creates a tiny sample document for each supported type and ingests it.
No external files needed.

Run from the project root:
    python scripts/test_ingestion.py
"""

import sys
import tempfile
from pathlib import Path

# Make sure project root is on the path
sys.path.insert(0, str(Path(__file__).parent.parent))

from utils.logging import setup_logging
from utils.settings import settings
from utils.models import DocType, ChunkType
from ingestion.pipeline import IngestionPipeline

setup_logging()

# ── colour helpers ────────────────────────────────────────────────────────
GREEN  = "\033[92m"
RED    = "\033[91m"
YELLOW = "\033[93m"
BOLD   = "\033[1m"
RESET  = "\033[0m"

def ok(msg):  print(f"{GREEN}  ✓ {msg}{RESET}")
def fail(msg): print(f"{RED}  ✗ {msg}{RESET}")
def info(msg): print(f"{YELLOW}  → {msg}{RESET}")


def make_txt(tmp: Path) -> Path:
    p = tmp / "sample.txt"
    p.write_text(
        "Annual Report 2023 — Acme Financial Services\n\n"
        "Revenue for the fiscal year ended 31 December 2023 was €12.4 million, "
        "representing a 14% increase over the prior year. EBITDA reached €3.1 million "
        "with a margin of 25%. Total debt stood at €8.2 million, giving a debt-to-equity "
        "ratio of 0.66. The company maintains a strong liquidity position with a current "
        "ratio of 2.1.\n\n"
        "Regional breakdown: Western Europe contributed 62% of revenue, Central and "
        "Eastern Europe 28%, and other markets the remaining 10%.\n\n"
        "Risk factors: Exposure to interest rate fluctuations, regulatory changes in "
        "the European Union, and currency risk from non-EUR denominated contracts."
        * 3  # repeat to get enough text for chunking
    )
    return p


def make_html(tmp: Path) -> Path:
    p = tmp / "sample.html"
    p.write_text("""<!DOCTYPE html>
<html>
<head><title>Financial Report</title></head>
<body>
  <nav>Navigation bar — should be stripped</nav>
  <h1>Q3 2023 Financial Results</h1>
  <p>Net revenue increased by 18% year-over-year to €4.2 million in Q3 2023.
     Operating costs were tightly managed, resulting in an operating margin of 22%.
     Cash and cash equivalents at quarter end were €1.8 million.</p>
  <h2>Segment Performance</h2>
  <p>The corporate lending segment grew 23% driven by new client acquisitions in
     the manufacturing and logistics sectors. Retail lending remained stable with
     a non-performing loan ratio of 1.8%, well below the industry average of 3.2%.</p>
  <footer>Footer content — should be stripped</footer>
</body>
</html>""")
    return p


def make_docx(tmp: Path) -> Path:
    """Create a minimal DOCX with paragraphs and a table."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return None  # skip if python-docx not installed

    doc = Document()
    doc.add_heading("Borrower Assessment Report", 0)
    doc.add_paragraph(
        "This report summarises the financial health of Meridian Logistics GmbH "
        "for the purpose of a €5 million term loan application. The company has "
        "demonstrated consistent revenue growth over the past three fiscal years "
        "and maintains adequate debt service coverage."
    )
    doc.add_heading("Key Financial Metrics", 1)

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    data = [
        ["Metric",            "FY2021",   "FY2022"],
        ["Revenue (€M)",      "8.2",      "9.7"],
        ["EBITDA Margin (%)", "19.4",     "21.1"],
        ["Debt / Equity",     "0.82",     "0.74"],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = val

    p = tmp / "sample.docx"
    doc.save(str(p))
    return p


# ── Test runner ────────────────────────────────────────────────────────────

def run_tests():
    pipe = IngestionPipeline(strategy="recursive", chunk_size=500, chunk_overlap=100)
    results = {}

    with tempfile.TemporaryDirectory() as tmp_str:
        tmp = Path(tmp_str)

        # ── TXT ──────────────────────────────────────────────
        print(f"\n{BOLD}[1] TXT ingestion{RESET}")
        try:
            f      = make_txt(tmp)
            chunks = pipe.ingest_file(f)
            assert len(chunks) > 0,            "Expected at least 1 chunk"
            assert all(c.text for c in chunks),"All chunks should have text"
            assert chunks[0].doc_id,           "doc_id should be set"
            assert chunks[0].token_count > 0,  "token_count should be > 0"
            ok(f"{len(chunks)} chunks | lang={chunks[0].language.value} "
               f"| first={chunks[0].text[:60].strip()!r}…")
            results["txt"] = True
        except Exception as e:
            fail(f"TXT failed: {e}")
            results["txt"] = False

        # ── HTML ─────────────────────────────────────────────
        print(f"\n{BOLD}[2] HTML ingestion{RESET}")
        try:
            f      = make_html(tmp)
            chunks = pipe.ingest_file(f)
            assert len(chunks) > 0
            # Make sure nav/footer was stripped
            full_text = " ".join(c.text for c in chunks).lower()
            assert "navigation bar" not in full_text, "Nav should have been stripped"
            assert "footer content" not in full_text, "Footer should have been stripped"
            assert "revenue"        in full_text,     "Content should be present"
            ok(f"{len(chunks)} chunks | nav/footer stripped ✓")
            results["html"] = True
        except Exception as e:
            fail(f"HTML failed: {e}")
            results["html"] = False

        # ── DOCX ─────────────────────────────────────────────
        print(f"\n{BOLD}[3] DOCX ingestion (+ table extraction){RESET}")
        f = make_docx(tmp)
        if f is None:
            info("python-docx not installed — DOCX test skipped")
            results["docx"] = None
        else:
            try:
                chunks = pipe.ingest_file(f)
                assert len(chunks) > 0
                text_chunks  = [c for c in chunks if c.chunk_type.value == "text"]
                table_chunks = [c for c in chunks if c.chunk_type.value == "table"]
                ok(f"{len(chunks)} total | {len(text_chunks)} text "
                   f"| {len(table_chunks)} table")
                if table_chunks:
                    ok(f"Table chunk preview:\n{table_chunks[0].text[:200]}")
                results["docx"] = True
            except Exception as e:
                fail(f"DOCX failed: {e}")
                results["docx"] = False

        # ── Directory batch ───────────────────────────────────
        print(f"\n{BOLD}[4] Directory batch ingestion{RESET}")
        try:
            all_results = pipe.ingest_directory(tmp, recursive=False)
            total_chunks = sum(len(v) for v in all_results.values())
            ok(f"{len(all_results)} files processed | {total_chunks} total chunks")
            for fname, ch in all_results.items():
                info(f"  {fname}: {len(ch)} chunks")
            results["directory"] = True
        except Exception as e:
            fail(f"Directory ingestion failed: {e}")
            results["directory"] = False

        # ── Chunker strategies ────────────────────────────────
        print(f"\n{BOLD}[5] Chunker strategy comparison{RESET}")
        from ingestion.chunker import DocumentChunker
        from utils.models import RawDocument, DocType

        dummy_doc  = RawDocument(source_path=Path("test.txt"), doc_type=DocType.TXT)
        dummy_text = "Financial analysis paragraph. " * 150

        for strat in ["recursive", "fixed"]:
            c = DocumentChunker(strategy=strat, chunk_size=400, chunk_overlap=80)
            ch = c.chunk(dummy_doc, dummy_text)
            ok(f"strategy={strat:<10} → {len(ch)} chunks, "
               f"avg_len={sum(len(x.text) for x in ch) // max(1,len(ch))} chars")

    # ── Summary ───────────────────────────────────────────────
    print(f"\n{BOLD}{'─'*50}")
    print("RESULTS")
    print('─'*50 + RESET)
    passed = skipped = failed = 0
    for name, result in results.items():
        if result is True:
            ok(name);     passed  += 1
        elif result is None:
            info(f"{name} (skipped — dependency missing)"); skipped += 1
        else:
            fail(name);   failed  += 1

    print(f"\n  {GREEN}{passed} passed{RESET}  "
          f"{YELLOW}{skipped} skipped{RESET}  "
          f"{RED}{failed} failed{RESET}\n")

    if failed:
        sys.exit(1)


if __name__ == "__main__":
    run_tests()