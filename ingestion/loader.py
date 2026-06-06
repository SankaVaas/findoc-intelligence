"""
Ingestion layer — loads raw files into (RawDocument, text, tables).

Supported:
  PDF  — digital text extraction (pypdf) + OCR fallback (Tesseract)
  DOCX — paragraphs + inline tables (python-docx)
  HTML — cleaned readable text (BeautifulSoup + html2text)
  TXT  — plain read
  Audio— transcript via Whisper (lazy-loaded, T4-optional)

Key design:
  - Every loader returns plain str; callers don't care how it was extracted
  - OCR is triggered per-page only when pypdf finds <20 chars (scanned pages)
  - Tables are extracted separately and returned alongside text
  - All errors are caught+logged; callers get "" rather than a crash
"""

from __future__ import annotations

import mimetypes
from pathlib import Path
from typing import Any

from utils.logging import logger
from utils.models import DocType, Language, RawDocument


# ─────────────────────────────────────────────────────────────
# Type detection
# ─────────────────────────────────────────────────────────────

def _detect_doc_type(path: Path) -> DocType:
    """Detect document type from magic bytes first, then extension."""
    try:
        import filetype
        kind = filetype.guess(str(path))
        if kind:
            m = kind.mime
            if "pdf"   in m: return DocType.PDF
            if "word"  in m or "officedocument" in m: return DocType.DOCX
            if "audio" in m or "video" in m: return DocType.AUDIO
    except ImportError:
        pass  # fall through to extension detection

    return {
        ".pdf":  DocType.PDF,
        ".docx": DocType.DOCX,
        ".doc":  DocType.DOCX,
        ".html": DocType.HTML,
        ".htm":  DocType.HTML,
        ".txt":  DocType.TXT,
        ".md":   DocType.TXT,
        ".mp3":  DocType.AUDIO,
        ".wav":  DocType.AUDIO,
        ".m4a":  DocType.AUDIO,
        ".ogg":  DocType.AUDIO,
    }.get(path.suffix.lower(), DocType.UNKNOWN)


# ─────────────────────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────────────────────

def _load_pdf_text(path: Path) -> str:
    """
    Extract text from PDF.
    Strategy:
      1. Try pypdf digital extraction per page
      2. If page has <20 chars → OCR that page with Tesseract
      3. If pypdf fails entirely → full OCR
    """
    try:
        import pypdf
    except ImportError:
        logger.error("pypdf not installed. Run: pip install pypdf")
        return ""

    parts: list[str] = []
    try:
        reader = pypdf.PdfReader(str(path))
        for page_num, page in enumerate(reader.pages):
            page_text = (page.extract_text() or "").strip()
            if len(page_text) < 20:
                logger.debug("Page {} looks scanned — trying OCR", page_num + 1)
                page_text = _ocr_single_page(path, page_num)
            parts.append(page_text)
        return "\n\n".join(parts)
    except Exception as e:
        logger.warning("pypdf failed ({}); attempting full OCR", e)
        return _ocr_full_pdf(path)


def _ocr_single_page(path: Path, page_num: int) -> str:
    """Rasterise one PDF page and run Tesseract on it."""
    try:
        import pytesseract
        from pdf2image import convert_from_path
        imgs = convert_from_path(str(path), first_page=page_num + 1,
                                 last_page=page_num + 1, dpi=300)
        if imgs:
            return pytesseract.image_to_string(imgs[0],
                                               lang="eng+fra+deu+spa+ita")
    except ImportError:
        logger.warning("pdf2image not installed — skipping OCR for page {}."
                       " Install: pip install pdf2image", page_num + 1)
    except Exception as e:
        logger.warning("OCR failed on page {}: {}", page_num + 1, e)
    return ""


def _ocr_full_pdf(path: Path) -> str:
    """OCR every page of a PDF (used when pypdf completely fails)."""
    try:
        import pytesseract
        from pdf2image import convert_from_path
        imgs = convert_from_path(str(path), dpi=300)
        return "\n\n".join(
            pytesseract.image_to_string(img, lang="eng+fra+deu+spa+ita")
            for img in imgs
        )
    except ImportError:
        logger.warning("pdf2image not installed — cannot OCR this PDF."
                       " Install: pip install pdf2image")
        return ""
    except Exception as e:
        logger.error("Full PDF OCR failed: {}", e)
        return ""


def _extract_pdf_tables(path: Path) -> list[dict[str, Any]]:
    """
    Extract tables from a PDF using pdfplumber.
    Returns list of {page, table_index, data: list[list[str|None]]}.
    """
    tables: list[dict[str, Any]] = []
    try:
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                for t_idx, table in enumerate(page.extract_tables() or []):
                    if table:
                        tables.append({
                            "page": page_num + 1,
                            "table_index": t_idx,
                            "data": table,
                        })
    except ImportError:
        logger.warning("pdfplumber not installed — table extraction skipped."
                       " Install: pip install pdfplumber")
    except Exception as e:
        logger.warning("PDF table extraction failed: {}", e)
    return tables


# ─────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────

def _load_docx(path: Path) -> tuple[str, list[dict[str, Any]]]:
    """
    Extract text and tables from a DOCX file.
    Returns (text, tables) where tables follows the same schema as PDF tables.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return "", []

    doc = DocxDocument(str(path))
    text_parts: list[str] = []
    tables: list[dict[str, Any]] = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    for t_idx, table in enumerate(doc.tables):
        rows = [
            [cell.text.strip() for cell in row.cells]
            for row in table.rows
        ]
        # Also add table content as text (for embedding)
        row_texts = [" | ".join(r) for r in rows if any(r)]
        text_parts.extend(row_texts)
        tables.append({"page": None, "table_index": t_idx, "data": rows})

    return "\n".join(text_parts), tables


# ─────────────────────────────────────────────────────────────
# HTML
# ─────────────────────────────────────────────────────────────

def _load_html(path: Path) -> str:
    try:
        from bs4 import BeautifulSoup
        import html2text
    except ImportError:
        logger.error("Install: pip install beautifulsoup4 lxml html2text")
        return path.read_text(encoding="utf-8", errors="replace")

    raw = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "lxml")

    # Strip boilerplate
    for tag in soup(["script", "style", "nav", "footer", "header",
                     "aside", "noscript", "iframe"]):
        tag.decompose()

    h = html2text.HTML2Text()
    h.ignore_links  = True
    h.ignore_images = True
    h.body_width    = 0
    return h.handle(str(soup)).strip()


# ─────────────────────────────────────────────────────────────
# Audio → transcript
# ─────────────────────────────────────────────────────────────

def _load_audio(path: Path, whisper_model: str = "base") -> str:
    """
    Transcribe audio using OpenAI Whisper (runs locally, free).
    Supports: MP3, WAV, M4A, OGG.
    On Colab T4 use whisper_model="medium" or "large" for accuracy.
    """
    try:
        import whisper
    except ImportError:
        logger.error("openai-whisper not installed. Run: pip install openai-whisper")
        return ""

    logger.info("Loading Whisper '{}' — this may take a minute on first run",
                whisper_model)
    try:
        model  = whisper.load_model(whisper_model)
        result = model.transcribe(str(path))
        text   = result["text"]
        lang   = result.get("language", "unknown")
        logger.info("Transcribed {} chars, detected language={}", len(text), lang)
        return text
    except Exception as e:
        logger.error("Whisper transcription failed: {}", e)
        return ""


# ─────────────────────────────────────────────────────────────
# Main loader class
# ─────────────────────────────────────────────────────────────

class DocumentLoader:
    """
    Load any supported document into (RawDocument, text, tables).

    Usage:
        loader = DocumentLoader()
        doc, text, tables = loader.load("report.pdf")

        # Batch
        results = loader.load_directory("data/raw/")
    """

    def __init__(self, whisper_model: str = "base"):
        self.whisper_model = whisper_model

    def load(
        self,
        path: str | Path,
        language_hint: Language | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> tuple[RawDocument, str, list[dict[str, Any]]]:
        """
        Load a single document.

        Returns:
            doc    — RawDocument with detected type, size, metadata
            text   — full extracted text (str)
            tables — list of table dicts extracted from the document
        """
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"Not found: {path}")

        doc_type = _detect_doc_type(path)
        logger.info("Loading '{}' as {}", path.name, doc_type.value)

        doc = RawDocument(
            source_path   = path,
            doc_type      = doc_type,
            language      = language_hint or Language.UNKNOWN,
            file_size_bytes = path.stat().st_size,
            metadata      = metadata or {},
        )

        text, tables = self._extract(path, doc_type)
        logger.info("Extracted {:,} chars, {} tables from '{}'",
                    len(text), len(tables), path.name)
        return doc, text, tables

    def _extract(
        self, path: Path, doc_type: DocType
    ) -> tuple[str, list[dict[str, Any]]]:
        """Dispatch to the right extractor and always return (text, tables)."""
        try:
            if doc_type == DocType.PDF:
                text   = _load_pdf_text(path)
                tables = _extract_pdf_tables(path)
                return text, tables

            if doc_type == DocType.DOCX:
                return _load_docx(path)          # already returns (text, tables)

            if doc_type == DocType.HTML:
                return _load_html(path), []

            if doc_type == DocType.TXT:
                return path.read_text(encoding="utf-8", errors="replace"), []

            if doc_type == DocType.AUDIO:
                return _load_audio(path, self.whisper_model), []

            logger.warning("No extractor for doc_type={}", doc_type)
            return "", []

        except Exception as e:
            logger.error("Extraction failed for '{}': {}", path.name, e)
            return "", []

    def load_directory(
        self,
        directory: str | Path,
        recursive: bool = True,
        extensions: list[str] | None = None,
    ) -> list[tuple[RawDocument, str, list[dict]]]:
        """
        Load all supported documents from a directory.
        Skips files that fail; logs errors per-file.
        """
        directory  = Path(directory)
        extensions = extensions or [".pdf", ".docx", ".doc",
                                    ".html", ".htm", ".txt",
                                    ".mp3", ".wav", ".m4a", ".ogg"]
        pattern = "**/*" if recursive else "*"
        results: list[tuple[RawDocument, str, list[dict]]] = []

        files = [
            f for f in directory.glob(pattern)
            if f.is_file() and f.suffix.lower() in extensions
        ]
        logger.info("Found {} files in '{}'", len(files), directory)

        for f in files:
            try:
                results.append(self.load(f))
            except Exception as e:
                logger.error("Skipping '{}': {}", f.name, e)

        logger.info("Successfully loaded {}/{} files", len(results), len(files))
        return results